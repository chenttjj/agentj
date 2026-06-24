#!/usr/bin/env python3
"""Build Mood Beat 專題報告.docx with embedded images."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = Path("report")
ASSETS = REPORT_DIR / "assets"
MD_PATH = REPORT_DIR / "專題報告.md"
OUT = REPORT_DIR / "專題報告.docx"

SAGE = RGBColor(0x87, 0x96, 0x73)
TERRACOTTA = RGBColor(0xC8, 0x6A, 0x4F)
DARK = RGBColor(0x33, 0x33, 0x33)

doc = Document()

# 標題
title = doc.add_heading("Mood Beat — 專題報告", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("馬公高中 Agent Studio 專題 — 完成日期 2026-06-24")
run.font.size = Pt(12)
run.font.color.rgb = DARK

# 1. 專題介紹
doc.add_heading("1. 專題介紹", level=1)
for k, v in [
    ("專題名稱", "Mood Beat"),
    ("一句話說明", "讓使用者描述今天狀態，透過呼吸穩定節奏，最後產生專屬 beat"),
    ("主要使用者", "喜歡音樂且喜歡心靈呼吸放鬆的人"),
    ("解決痛點", "壓力大、心情不愉快"),
]:
    p = doc.add_paragraph()
    p.add_run(f"{k}：").bold = True
    p.add_run(v)

# 2. 學校 Server 環境
doc.add_heading("2. 學校 Server 環境", level=1)
p = doc.add_paragraph()
p.add_run("全班使用相同的學校內部 LLM 路由，所有 API 連線皆透過此路由介接。本報告不揭露 IP／URL 細節。")
doc.add_picture(str(ASSETS / "server-topology.png"), width=Inches(6.0))
last = doc.paragraphs[-1]
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 3. 左欄與右欄互動
doc.add_heading("3. 左欄與右欄互動說明", level=1)
doc.add_heading("3.1 左欄自訂頁", level=2)
for k, v in [
    ("4_Layout_Practice.py", "版面練習（demo 參考頁）"),
    ("5_Mood_Beat.py", "核心頁：輸入情緒、啟動呼吸、生成 beat"),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{k}").bold = True
    p.add_run(f" — {v}")

doc.add_heading("3.2 Mood Beat 頁面輸入元件", level=2)
for item in [
    "心情值滑桿（1–10）",
    "壓力值滑桿（1–10）",
    "音樂風格下拉選單",
    "「開始呼吸」按鈕",
    "「生成 Beat」按鈕",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3.3 資料傳遞", level=2)
p = doc.add_paragraph()
p.add_run("左欄把「心情值、壓力值、音樂偏好」三項資料整理成一段文字摘要，傳給右欄 Agent。Agent 推理後回傳：")
for item in [
    "文字建議（呼吸節奏提示、推薦音樂風格）",
    "JSON 節拍參數（含節拍 BPM、情緒標籤）",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3.4 完整互動例子", level=2)
p = doc.add_paragraph()
p.add_run("我按了「生成 Beat」按鈕，Agent 收到心情值 7、壓力值 3、音樂偏好 lo-fi，然後左欄顯示呼吸節奏提示和「你的專屬 lo-fi beat 已生成」。").italic = True

doc.add_heading("3.5 個人架構圖", level=2)
doc.add_picture(str(ASSETS / "project-architecture.png"), width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 4. 成果、創新與技術
doc.add_heading("4. 成果、創新與技術", level=1)
doc.add_heading("4.1 成果", level=2)
for item in [
    "在左欄 Mood Beat 頁面用滑桿輸入心情值與壓力值",
    "選擇喜歡的音樂風格（如 lo-fi、trap、輕音樂）",
    "按下「開始呼吸」按鈕啟動呼吸引導",
    "按下「生成 Beat」按鈕，將情緒資料傳給右欄 Agent",
    "右欄 Agent 回傳文字建議與 JSON 節拍參數，左欄顯示推薦結果",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.2 創新／亮點", level=2)
for item in [
    "結合「情緒紀錄」「呼吸穩定」「音樂推薦」三個步驟，不只是純聊天或純音樂播放",
    "用心情值、壓力值、音樂偏好三項資料，讓 Agent 產出個人化的 beat 風格建議",
    "強調「先穩定心情、再創作音樂」的體驗流程，幫助使用者紓壓",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.3 技術含量", level=2)
for item in [
    "Agent Studio + Streamlit 打造左欄自訂頁面",
    "Agent.chat 作為右欄對話與推理核心",
    "JSON 格式傳遞心情值、壓力值、音樂偏好",
    "LLM 根據情緒資料產生呼吸建議與音樂風格推薦",
    "未來可串接 ACE Music 等音樂生成工具輸出實際 beat",
]:
    doc.add_paragraph(item, style="List Bullet")

# 5. Demo 截圖
doc.add_heading("5. 附錄：Demo 截圖", level=1)
doc.add_paragraph("demo-02：Mood Beat 頁面實際畫面", style="List Bullet")
doc.add_picture(str(ASSETS / "demo-02.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f"OK: {OUT}")
